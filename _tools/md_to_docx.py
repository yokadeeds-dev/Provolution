"""
md_to_docx.py — Convert Provolution manuscript markdown to .docx
Usage: python md_to_docx.py <input.md> <output.docx>
"""
import sys
import re
from docx import Document
from docx.shared import Pt

def add_inline(para, text):
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = para.add_run(part[3:-3])
            run.bold = True; run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
        else:
            para.add_run(part)

def convert(src, dst):
    doc = Document()
    # Base font
    for style in doc.styles:
        try:
            style.font.name = 'Calibri'
        except Exception:
            pass

    with open(src, encoding='utf-8') as f:
        lines = f.readlines()

    in_code = False
    code_lines = []
    table_rows = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # ── Code block ──────────────────────────────────────────
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                if code_lines:
                    p = doc.add_paragraph('\n'.join(code_lines))
                    p.style = 'No Spacing'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Table ────────────────────────────────────────────────
        if line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            # separator row
            if all(re.fullmatch(r'[-: ]+', c) for c in cells if c):
                i += 1
                continue
            table_rows.append(cells)
            # flush when table ends
            next_is_table = (i + 1 < len(lines) and lines[i+1].strip().startswith('|'))
            if not next_is_table and table_rows:
                ncols = max(len(r) for r in table_rows)
                t = doc.add_table(rows=len(table_rows), cols=ncols)
                t.style = 'Table Grid'
                for ri, row in enumerate(table_rows):
                    for ci in range(ncols):
                        cell_text = row[ci] if ci < len(row) else ''
                        cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                        cell_text = re.sub(r'\*(.+?)\*', r'\1', cell_text)
                        t.rows[ri].cells[ci].text = cell_text
                        if ri == 0:
                            for run in t.rows[ri].cells[ci].paragraphs[0].runs:
                                run.bold = True
                table_rows = []
                doc.add_paragraph()
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────
        m = re.match(r'^(#{1,4}) (.+)', line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────
        if re.fullmatch(r'-{3,}', line.strip()):
            p = doc.add_paragraph('─' * 50)
            i += 1
            continue

        # ── Bullet list ──────────────────────────────────────────
        if re.match(r'^[-*] ', line):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, line[2:])
            i += 1
            continue

        # ── Numbered list ────────────────────────────────────────
        if re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            add_inline(p, re.sub(r'^\d+\. ', '', line))
            i += 1
            continue

        # ── Empty line ───────────────────────────────────────────
        if line.strip() == '':
            doc.add_paragraph()
            i += 1
            continue

        # ── Normal paragraph ─────────────────────────────────────
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    doc.save(dst)
    print(f'Saved: {dst}')

if __name__ == '__main__':
    src = sys.argv[1] if len(sys.argv) > 1 else 'MANUSCRIPT_DRAFT_v0.1.md'
    dst = sys.argv[2] if len(sys.argv) > 2 else src.replace('.md', '.docx')
    convert(src, dst)
